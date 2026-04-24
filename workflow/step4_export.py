"""
Step 4: 导出为 Word 文档
"""
import re
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from rich.console import Console

console = Console()


def _set_font(run, font_name: str = "微软雅黑", size: int = 12, bold: bool = False, color: tuple = None):
    """设置字体样式"""
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if color:
        run.font.color.rgb = RGBColor(*color)


def export_to_word(
    content: str,
    topic: str,
    analysis: str = "",
    output_dir: str = None,
    filename: str = None,
) -> str:
    """
    将文案导出为 Word 文档

    Args:
        content: 文案内容（含正式稿和注释版）
        topic: 创作选题
        analysis: 爆款分析报告
        output_dir: 输出目录
        filename: 文件名（可选）

    Returns:
        保存的文件路径
    """
    if output_dir is None:
        from config import get_preferences
        prefs = get_preferences()
        output_dir = prefs.get("output_dir", str(Path.home() / "Desktop"))

    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)

    if filename is None:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M")
        safe_topic = re.sub(r'[\\/:*?"<>|]', '_', topic)[:30]
        filename = f"文案_{safe_topic}_{timestamp}.docx"

    file_path = output_path / filename

    doc = Document()

    # 页面设置
    for section in doc.sections:
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)

    # === 封面信息 ===
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(f"📝 {topic}")
    _set_font(title_run, size=20, bold=True, color=(31, 73, 125))

    doc.add_paragraph()
    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta_para.add_run(
        f"生成时间：{datetime.now().strftime('%Y年%m月%d日 %H:%M')}  |  AI 视频文案助手"
    )
    _set_font(meta_run, size=10, color=(128, 128, 128))

    doc.add_paragraph()
    _add_divider(doc)
    doc.add_paragraph()

    # === 解析内容块 ===
    sections = _parse_content_sections(content)

    # 正式稿
    if "正式稿" in sections:
        _add_section_header(doc, "🎬 正式稿（可直接录制）")
        _add_body_text(doc, sections["正式稿"])
        doc.add_paragraph()

    # 注释版
    if "注释版" in sections:
        _add_divider(doc)
        doc.add_paragraph()
        _add_section_header(doc, "📋 注释版（含结构标注）", color=(0, 102, 0))
        _add_annotated_text(doc, sections["注释版"])
        doc.add_paragraph()

    # 创作说明
    if "创作说明" in sections:
        _add_divider(doc)
        doc.add_paragraph()
        _add_section_header(doc, "💡 创作说明", color=(102, 0, 102))
        _add_body_text(doc, sections["创作说明"])
        doc.add_paragraph()

    # 爆款分析报告（如果有）
    if analysis:
        _add_divider(doc)
        doc.add_paragraph()
        _add_section_header(doc, "📊 参考爆款分析报告", color=(102, 51, 0))
        _add_body_text(doc, analysis[:3000] + ("..." if len(analysis) > 3000 else ""))

    doc.save(str(file_path))
    console.print(f"\n[bold green]✅ 文档已保存：{file_path}[/bold green]")
    return str(file_path)


def _parse_content_sections(content: str) -> dict:
    """解析文案内容中的各个部分"""
    sections = {}
    # 匹配 # 【XXX】 格式的标题
    pattern = r'#\s*【([^】]+)】\s*\n(.*?)(?=\n#\s*【|\Z)'
    matches = re.findall(pattern, content, re.DOTALL)
    for title, body in matches:
        sections[title.strip()] = body.strip()

    # 如果没有匹配到标准格式，把全部内容作为正式稿
    if not sections:
        sections["正式稿"] = content.strip()

    return sections


def _add_section_header(doc, text: str, color: tuple = (31, 73, 125)):
    """添加章节标题"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    _set_font(run, size=14, bold=True, color=color)


def _add_body_text(doc, text: str):
    """添加正文内容"""
    paragraphs = text.split("\n")
    for line in paragraphs:
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
        # 检测是否是 Markdown 标题
        if line.startswith("## "):
            para = doc.add_paragraph()
            run = para.add_run(line[3:])
            _set_font(run, size=13, bold=True, color=(31, 73, 125))
        elif line.startswith("### "):
            para = doc.add_paragraph()
            run = para.add_run(line[4:])
            _set_font(run, size=12, bold=True, color=(0, 70, 127))
        elif line.startswith("- ") or line.startswith("• "):
            para = doc.add_paragraph(style='List Bullet')
            run = para.add_run(line[2:])
            _set_font(run, size=11)
        else:
            para = doc.add_paragraph()
            run = para.add_run(line)
            _set_font(run, size=12)
            para.paragraph_format.space_after = Pt(4)
            para.paragraph_format.line_spacing = Pt(22)


def _add_annotated_text(doc, text: str):
    """添加带注释标注的文本（将 [标注] 高亮显示）"""
    paragraphs = text.split("\n")
    for line in paragraphs:
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue

        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(4)
        para.paragraph_format.line_spacing = Pt(22)

        # 检测注释标记 [xxx]
        annotation_pattern = re.compile(r'(\[[^\]]+\])')
        parts = annotation_pattern.split(line)

        for part in parts:
            if annotation_pattern.match(part):
                # 这是注释标记，用彩色显示
                run = para.add_run(part)
                _set_font(run, size=10, color=(255, 102, 0))
                run.font.bold = True
            else:
                run = para.add_run(part)
                _set_font(run, size=12)


def _add_divider(doc):
    """添加分割线"""
    para = doc.add_paragraph("─" * 50)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in para.runs:
        run.font.color.rgb = RGBColor(200, 200, 200)
